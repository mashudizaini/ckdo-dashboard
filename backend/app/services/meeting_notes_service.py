"""
Meeting Notes Service
─────────────────────────────────────────
Transcription runs on the remote GPU Whisper service on the "ai-engine" VM
(172.21.2.27:9500 — faster-whisper large-v3, systemd unit
whisper-server.service) instead of running faster-whisper in-process on the
backend's CPU — ~17x realtime on that box's RTX 5060 Ti, so a 1-2 hour
meeting transcribes in a few minutes instead of potentially longer than the
meeting itself. See app/config.py's whisper_api_url comment. Transcription
itself has no Claude alternative — the Anthropic API has no audio input
capability — so it always runs on-premise regardless of provider choice.

MOM generation supports 6 providers: "onprem" (default, local Ollama —
see config.py's ollama_mom_model — free), "anthropic" (Claude), "gemini"
(reuses gemini_service.py), "deepseek", "openai" (ChatGPT), or "kimi"
(Moonshot AI) — the last 3 are raw REST, OpenAI-compatible endpoints (see
app/config.py's deepseek/openai/kimi settings). All six return the same
{"departments": [...]} JSON shape from the same MOM_PROMPT_TEMPLATE, so
quality differences are purely down to the underlying model, not the
instructions given to it.

Each cloud provider (anthropic/gemini/deepseek/openai/kimi) will use the
calling user's own saved key (see user_api_key_service.py) when one
exists, falling back to the shared company key from Settings otherwise —
resolved by the router before calling generate_mom(), passed in via the
api_key parameter.

MOM schema (departments -> topics -> discussion_points + action_plans) and
the rendered DOCX layout match the company's previous meeting-notes tool
(sumber/CKDO_DASHBOARD/apps/meeting_notes_app) for continuity with the
format staff are already used to — see build_mom_docx().
"""
import json
import re
import io
import httpx
import anthropic
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.config import get_settings
from app.services import gemini_service
import structlog

logger = structlog.get_logger()
settings = get_settings()


class MomProviderCreditError(Exception):
    """Raised when a cloud MOM provider (Anthropic/Gemini/DeepSeek) rejects
    the request because the account is out of credit/quota — a distinct,
    actionable case the router surfaces as its own clear message instead of
    a generic 500 with the provider's raw (often cryptic) error body."""


def _is_credit_error(status_code: int, body: str) -> bool:
    if status_code not in (400, 402, 403, 429):
        return False
    return bool(re.search(r"insufficient|credit balance|quota|resource_exhausted|payment required", body, re.I))

OLLAMA_MOM_TIMEOUT_SECONDS = 120.0
CLOUD_MOM_TIMEOUT_SECONDS = 180.0
# Cloud providers can run genuinely long, detailed MOMs — this is a JSON
# structure with many departments/topics/points, not a short chat reply.
MOM_MAX_OUTPUT_TOKENS = 8192

# Generous margin — ~17x realtime measured on the ai-engine GPU means even a
# 5h meeting (the longest meetings are expected to run) takes under 20
# minutes to transcribe, but a large upload + network hop deserves headroom
# over the bare compute time. nginx's proxy_*_timeout (nginx.dev.conf) is
# set to clear this with its own margin — raise both together.
TRANSCRIBE_TIMEOUT_SECONDS = 3600.0

# Ollama supports grammar-constrained decoding via `format: <json schema>` —
# the sampler is restricted to only emit tokens that keep the output valid
# against this shape, so nesting mistakes (a topic dropped a level, a missing
# array close before the next department) become structurally impossible
# instead of a ~30-50% failure rate. Measured empirically on this deployment
# (qwen2.5:14b-instruct and qwen3:14b, 4 runs each): 3/6 malformed without
# this constraint, 8/8 valid with it — and each run was also faster, since
# the model doesn't waste tokens rambling once state is pinned down. Only
# applied to the on-prem branch: the cloud providers (Claude/Gemini/DeepSeek)
# are frontier models that already keep this schema straight reliably.
MOM_JSON_SCHEMA = {
    "type": "object",
    "properties": {
        "departments": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "topics": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "title": {"type": "string"},
                                "discussion_points": {"type": "array", "items": {"type": "string"}},
                                "action_plans": {"type": "array", "items": {"type": "string"}},
                            },
                            "required": ["title", "discussion_points", "action_plans"],
                        },
                    },
                },
                "required": ["name", "topics"],
            },
        },
    },
    "required": ["departments"],
}

MOM_PROMPT_TEMPLATE = """Analisis transkrip rapat berikut dan buatkan ringkasan terstruktur dalam format JSON.

{meta}Transkrip rapat:
\"\"\"
{transcript}
\"\"\"

ATURAN:
- Organisasikan berdasarkan departemen/topik utama yang benar-benar dibahas dalam transkrip (misal: HRGA, IT, Purchasing, Accounting & Tax, Planning & Coordination, dst — jangan mengarang departemen yang tidak disebut). Tulis nama departemen dengan kapitalisasi wajar (akronim seperti HRGA/IT/PAC dalam huruf besar, nama lain dalam Title Case) — jangan menuliskan semua nama departemen dalam huruf besar semua.
- TANGKAP DETAIL SELENGKAP MUNGKIN. Jangan meringkas satu topik yang dibahas panjang jadi satu kalimat umum — pecah jadi beberapa discussion_points terpisah. Kalau transkrip menyebut angka, nama orang/vendor/klien, alasan/latar belakang, atau rincian teknis, sertakan itu apa adanya di discussion_points, bukan cuma judul topiknya secara umum.
- Setiap topik punya discussion_points (poin pembahasan) dan action_plans (rencana tindakan untuk minggu berikutnya). Untuk action_plans, sertakan PIC dan deadline jika disebutkan, format: "tindakan - oleh Nama (deadline)". Kalau tidak ada rencana tindakan untuk topik itu, kembalikan array kosong.
- Setiap poin di discussion_points/action_plans sebaiknya diawali label singkat dalam **bold** (mis. "**Recruitment**: ...", "**Annual MCU**: ...") kalau poin itu membahas sub-topik/item yang jelas — ini mengikuti gaya laporan MOM resmi perusahaan.
- Gunakan bahasa formal, abaikan basa-basi/obrolan yang benar-benar tidak substantif (candaan, sapaan) — tapi JANGAN buang detail substantif hanya demi keringkasan.
- SELALU tulis seluruh isi JSON (nama departemen, judul topik, discussion_points, action_plans) dalam BAHASA INGGRIS, apa pun bahasa transkripnya — kalau transkrip berbahasa Indonesia, terjemahkan ke Bahasa Inggris profesional untuk laporan ini.
- Jangan mengarang informasi yang tidak ada di transkrip.

Balas HANYA dengan JSON valid (tanpa markdown code fence, tanpa teks lain), struktur PERSIS seperti ini:
{{
  "departments": [
    {{
      "name": "Department Name",
      "topics": [
        {{
          "title": "Topic title",
          "discussion_points": ["**Label**: point 1", "point 2"],
          "action_plans": ["action - by Name (deadline)"]
        }}
      ]
    }}
  ]
}}
"""


def _add_markdown_runs(paragraph, text: str):
    """Split text on **bold** markers into runs — same lightweight Markdown
    convention used for the PPTX exports (business_plan_setup_service.py),
    applied here to docx runs instead of pptx runs."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        run.bold = is_bold


def _add_horizontal_line(paragraph):
    """Word has no native "horizontal rule" — a bottom paragraph border is
    the standard docx trick to render one."""
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(p_bdr)


class MeetingNotesService:

    async def transcribe(self, file_bytes: bytes, filename: str, language: str | None = None) -> dict:
        """Upload audio to the remote GPU Whisper service, return the
        transcript + segments. Raises httpx.HTTPStatusError / TimeoutException
        on failure — the router translates these into a clean error response."""
        url = f"{settings.whisper_api_url.rstrip('/')}/transcribe"
        files = {"file": (filename, file_bytes)}
        data = {"language": language} if language else {}
        async with httpx.AsyncClient(timeout=TRANSCRIBE_TIMEOUT_SECONDS) as client:
            resp = await client.post(url, files=files, data=data)
            resp.raise_for_status()
            return resp.json()

    def extract_text_from_upload(self, file_bytes: bytes, filename: str) -> str:
        """Pulls plain text out of an uploaded transcript file — .docx via
        python-docx (paragraph text only, same library already used for
        build_mom_docx below), anything else decoded as plain UTF-8 text
        (.txt/.srt/.vtt/etc.). Raises ValueError with a user-facing message
        on an unreadable/corrupt file."""
        ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
        if ext == "docx":
            try:
                doc = Document(io.BytesIO(file_bytes))
            except Exception as e:
                raise ValueError(f"Tidak bisa membaca file .docx: {e}") from e
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Tidak bisa membaca file sebagai teks: {e}") from e

    async def generate_mom(
        self, transcript: str, meeting_title: str = "", participants: str = "",
        provider: str = "onprem", api_key: str | None = None,
    ) -> dict:
        """Transcript -> {"departments": [...]}. provider: "onprem" (default,
        local Ollama), "anthropic" (Claude), "gemini", "deepseek", "openai"
        (ChatGPT), or "kimi" (Moonshot AI). api_key, when given, overrides
        the shared company key for that provider (resolved by the router
        from the calling user's own saved key, if any) — ignored for
        "onprem" since that's a free local model with no key at all."""
        meta_lines = []
        if meeting_title:
            meta_lines.append(f"Judul rapat: {meeting_title}")
        if participants:
            meta_lines.append(f"Peserta: {participants}")
        meta = ("\n".join(meta_lines) + "\n\n") if meta_lines else ""

        prompt = MOM_PROMPT_TEMPLATE.format(meta=meta, transcript=transcript)

        if provider == "anthropic":
            try:
                client = anthropic.AsyncAnthropic(api_key=api_key or settings.anthropic_api_key)
                response = await client.messages.create(
                    model="claude-opus-5",
                    max_tokens=MOM_MAX_OUTPUT_TOKENS,
                    messages=[{"role": "user", "content": prompt}],
                )
            except anthropic.APIStatusError as e:
                if _is_credit_error(e.status_code, str(e)):
                    raise MomProviderCreditError(
                        "Saldo/kredit API Claude (Anthropic) tidak cukup. Silakan isi ulang di "
                        "console.anthropic.com > Plans & Billing, atau pilih provider lain."
                    ) from e
                raise
            raw = response.content[0].text.strip()
        elif provider == "gemini":
            try:
                raw = (await gemini_service.generate(
                    system_prompt="You produce structured meeting minutes. Respond with valid JSON only, no commentary.",
                    contents=[{"role": "user", "parts": [{"text": prompt}]}],
                    api_key=api_key,
                )).strip()
            except httpx.HTTPStatusError as e:
                if _is_credit_error(e.response.status_code, e.response.text):
                    raise MomProviderCreditError(
                        "Kuota/saldo API Gemini tidak cukup. Silakan cek kuota di Google AI Studio, "
                        "atau pilih provider lain."
                    ) from e
                raise
        elif provider == "deepseek":
            async with httpx.AsyncClient(timeout=CLOUD_MOM_TIMEOUT_SECONDS) as client:
                resp = await client.post(
                    "https://api.deepseek.com/chat/completions",
                    headers={"Authorization": f"Bearer {api_key or settings.deepseek_api_key}"},
                    json={
                        "model": settings.deepseek_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": MOM_MAX_OUTPUT_TOKENS,
                        "stream": False,
                    },
                )
                if _is_credit_error(resp.status_code, resp.text):
                    raise MomProviderCreditError(
                        "Saldo API DeepSeek tidak cukup. Silakan top up di platform.deepseek.com, "
                        "atau pilih provider lain."
                    )
                resp.raise_for_status()
                raw = resp.json()["choices"][0]["message"]["content"].strip()
        elif provider == "openai":
            async with httpx.AsyncClient(timeout=CLOUD_MOM_TIMEOUT_SECONDS) as client:
                resp = await client.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={"Authorization": f"Bearer {api_key or settings.openai_api_key}"},
                    json={
                        "model": settings.openai_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": MOM_MAX_OUTPUT_TOKENS,
                        "response_format": {"type": "json_object"},
                        "stream": False,
                    },
                )
                if _is_credit_error(resp.status_code, resp.text):
                    raise MomProviderCreditError(
                        "Saldo/kredit API OpenAI (ChatGPT) tidak cukup. Silakan cek billing di "
                        "platform.openai.com, atau pilih provider lain."
                    )
                resp.raise_for_status()
                raw = resp.json()["choices"][0]["message"]["content"].strip()
        elif provider == "kimi":
            async with httpx.AsyncClient(timeout=CLOUD_MOM_TIMEOUT_SECONDS) as client:
                resp = await client.post(
                    f"{settings.kimi_api_base.rstrip('/')}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key or settings.kimi_api_key}"},
                    json={
                        "model": settings.kimi_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": MOM_MAX_OUTPUT_TOKENS,
                        "stream": False,
                    },
                )
                if _is_credit_error(resp.status_code, resp.text):
                    raise MomProviderCreditError(
                        "Saldo API Kimi (Moonshot AI) tidak cukup. Silakan top up di platform.moonshot.ai, "
                        "atau pilih provider lain."
                    )
                resp.raise_for_status()
                raw = resp.json()["choices"][0]["message"]["content"].strip()
        else:
            async with httpx.AsyncClient(timeout=OLLAMA_MOM_TIMEOUT_SECONDS) as client:
                resp = await client.post(
                    f"{settings.ollama_api_url.rstrip('/')}/api/chat",
                    json={
                        "model": settings.ollama_mom_model,
                        "messages": [{"role": "user", "content": prompt}],
                        "stream": False,
                        # Ollama's default temperature (~0.8) is tuned for open-ended
                        # chat, not structured extraction — for a task like "pull the
                        # exact points out of this transcript" that randomness mostly
                        # shows up as inconsistency/invention rather than useful
                        # variety. Low temperature + constrained top_p keep the smaller
                        # on-prem model closer to the transcript's actual content.
                        "options": {"temperature": 0.15, "top_p": 0.9},
                        "format": MOM_JSON_SCHEMA,
                        # Harmless no-op for non-thinking models (confirmed empirically);
                        # required for qwen3-class hybrid-thinking models so the chain-
                        # of-thought preamble doesn't leak into `message.content` ahead
                        # of the JSON.
                        "think": False,
                    },
                )
                resp.raise_for_status()
                raw = resp.json()["message"]["content"].strip()

        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)

    def build_mom_docx(self, mom_json: dict, meeting_title: str, participants: str, meta: dict) -> bytes:
        """Renders the (possibly user-edited) MOM structure into a .docx,
        matching the company's actual MOM template 1:1 (sumber/4. MOM Admin
        Jul 24, 2026.pdf): a letterhead-style two-column header (company
        name left, "Minutes of Meeting - {title}" right), Date/Time/Venue,
        a column-major two-column participants list, Agenda, a horizontal
        rule, "[Discussed matters]", then per-department bold+underlined
        headers with numbered topics, "o Discussion Points" / "o Next
        Week/Action Plan" sub-headers, and "▪"-bulleted points (with inline
        **bold** labels rendered as real bold runs)."""
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        usable_width = section.page_width - section.left_margin - section.right_margin

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # ── Letterhead header: company name left, "Minutes of Meeting - X" right ──
        header = doc.add_paragraph()
        header.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
        header.add_run("PT CKD OTTO Pharmaceuticals").bold = True
        header.add_run("\t")
        header.add_run(f"Minutes of Meeting - {meeting_title or 'Meeting'}").bold = True
        doc.add_paragraph()

        doc.add_paragraph(f"Date : {meta.get('date', '')}")
        doc.add_paragraph(f"Time : {meta.get('time', '')}")
        doc.add_paragraph(f"Venue : {meta.get('venue', '')}")
        doc.add_paragraph()

        names = [n.strip() for n in (participants or "").split(",") if n.strip()]
        if names:
            doc.add_paragraph("Participants:")
            rows = (len(names) + 1) // 2
            table = doc.add_table(rows=rows, cols=2)
            for i, name in enumerate(names):
                # Column-major fill (1-4 down the left column, 5-7 down the
                # right), matching the reference — not row-major left-right.
                col, row = i // rows, i % rows
                table.cell(row, col).text = f"{i + 1}. {name}"
            doc.add_paragraph()

        if meta.get("agenda"):
            doc.add_paragraph(f"Agenda: {meta['agenda']}")

        rule_p = doc.add_paragraph()
        _add_horizontal_line(rule_p)

        marker = doc.add_paragraph()
        marker.add_run("[Discussed matters]").bold = True
        doc.add_paragraph()

        for dept in mom_json.get("departments", []):
            dept_p = doc.add_paragraph()
            dept_run = dept_p.add_run(dept.get("name") or "")
            dept_run.bold = True
            dept_run.underline = True

            for i, topic in enumerate(dept.get("topics", []), start=1):
                topic_p = doc.add_paragraph()
                topic_p.paragraph_format.left_indent = Cm(0.5)
                topic_p.add_run(f"{i}. {topic.get('title', '')}").bold = True

                points = topic.get("discussion_points") or []
                if points:
                    sub = doc.add_paragraph()
                    sub.paragraph_format.left_indent = Cm(1.0)
                    sub.add_run("o Discussion Points").bold = True
                    for point in points:
                        bp = doc.add_paragraph()
                        bp.paragraph_format.left_indent = Cm(1.5)
                        bp.add_run("▪ ")
                        _add_markdown_runs(bp, point)

                actions = topic.get("action_plans") or []
                if actions:
                    sub = doc.add_paragraph()
                    sub.paragraph_format.left_indent = Cm(1.0)
                    sub.add_run("o Next Week/Action Plan").bold = True
                    for action in actions:
                        bp = doc.add_paragraph()
                        bp.paragraph_format.left_indent = Cm(1.5)
                        bp.add_run("▪ ")
                        _add_markdown_runs(bp, action)

            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
