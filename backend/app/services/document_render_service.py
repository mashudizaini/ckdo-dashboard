"""
Document Converter — Rendering Service
─────────────────────────────────────────
Renders the structured block list (document_converter_service.extract_
blocks / document_translation_service.translate_blocks) into MD/DOCX/XLSX/
JSONL. Every format renders from that SAME block list, not from each
other's output, so a table cell is never re-parsed out of another
format's flattened text (the failure mode
sumber/Panduan_Konversi_dan_Terjemahan_Dokumen_KO.md opens with).

Bilingual ("both") mode renders the original section followed by the
translated section, each with its own real table — not merged cell-by-
cell, which would need row/column alignment guarantees this app's
generic docling-based extraction doesn't provide (see the deferred
Phase-2 coordinate-extraction work noted in the enhancement plan).
"""
import io
import json

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from fastapi.responses import StreamingResponse

from app.services.document_converter_service import render_markdown_from_blocks

TITLE_FONT  = Font(bold=True, size=14)
HEADER_FONT = Font(bold=True, size=10, color="FFFFFF")
HEADER_FILL = PatternFill("solid", fgColor="1F4E78")
LEFT_WRAP   = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _stream(buf: io.BytesIO, fname: str, media_type: str) -> StreamingResponse:
    buf.seek(0)
    return StreamingResponse(buf, media_type=media_type, headers={"Content-Disposition": f"attachment; filename={fname}"})


def _max_table_cols(blocks: list) -> int:
    widths = [len(row) for b in blocks for row in b.get("rows", []) if b.get("type") == "table"]
    return max(widths, default=1)


# ── Markdown ─────────────────────────────────────────────────────────────

def render_md_response(blocks: list, blocks_translated: list | None, title: str, fname: str) -> StreamingResponse:
    parts = [f"# {title}"] if title else []
    parts.append(render_markdown_from_blocks(blocks))
    if blocks_translated:
        parts.append("---\n\n## Translated")
        parts.append(render_markdown_from_blocks(blocks_translated))
    buf = io.BytesIO("\n\n".join(p for p in parts if p).encode("utf-8"))
    return _stream(buf, fname, "text/markdown; charset=utf-8")


# ── DOCX — real tables via doc.add_table(), never markdown pipe-text ─────

def _write_docx_section(doc: DocxDocument, blocks: list, section_title: str | None = None):
    if section_title:
        doc.add_heading(section_title, level=1)
    for b in blocks:
        if b.get("type") == "heading":
            doc.add_heading(b.get("text", ""), level=min(max(b.get("level", 1), 1), 4) + 1)
        elif b.get("type") == "paragraph":
            if b.get("text"):
                doc.add_paragraph(b["text"])
        elif b.get("type") == "table" and b.get("rows"):
            header, *body = b["rows"]
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Table Grid"
            for cell, text in zip(t.rows[0].cells, header):
                cell.paragraphs[0].add_run(text or "").bold = True
            for r in body:
                cells = t.add_row().cells
                for i, val in enumerate(r[:len(header)]):
                    cells[i].text = val or ""
            doc.add_paragraph("")


def render_docx_response(blocks: list, blocks_translated: list | None, title: str, fname: str) -> StreamingResponse:
    doc = DocxDocument()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading(title or "Document", level=0)
    if blocks_translated:
        _write_docx_section(doc, blocks, section_title="Original")
        _write_docx_section(doc, blocks_translated, section_title="Translated")
    else:
        _write_docx_section(doc, blocks)

    buf = io.BytesIO()
    doc.save(buf)
    return _stream(buf, fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ── XLSX — real cells via openpyxl, styled to match this app's other
#           exports (financial_statement.py's title/header conventions) ──

def _write_xlsx_blocks(ws, blocks: list, start_row: int) -> int:
    row = start_row
    for b in blocks:
        if b.get("type") == "heading":
            c = ws.cell(row=row, column=1, value=b.get("text", ""))
            c.font = Font(bold=True, size=12)
            row += 1
        elif b.get("type") == "paragraph":
            if b.get("text"):
                ws.cell(row=row, column=1, value=b["text"])
                row += 1
        elif b.get("type") == "table" and b.get("rows"):
            header, *body = b["rows"]
            for ci, text in enumerate(header, start=1):
                c = ws.cell(row=row, column=ci, value=text)
                c.font, c.fill, c.alignment = HEADER_FONT, HEADER_FILL, LEFT_WRAP
            row += 1
            for r in body:
                for ci, val in enumerate(r[:len(header)], start=1):
                    ws.cell(row=row, column=ci, value=val)
                row += 1
            row += 1  # spacer between tables
    return row


def _size_columns(ws, ncols: int):
    ws.column_dimensions["A"].width = 50
    for i in range(2, ncols + 1):
        ws.column_dimensions[get_column_letter(i)].width = 22


def render_xlsx_response(blocks: list, blocks_translated: list | None, title: str, fname: str) -> StreamingResponse:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Original"
    ws.cell(row=1, column=1, value=title or "Document").font = TITLE_FONT
    _write_xlsx_blocks(ws, blocks, start_row=3)
    _size_columns(ws, _max_table_cols(blocks))

    if blocks_translated:
        ws2 = wb.create_sheet("Translated")
        ws2.cell(row=1, column=1, value=title or "Document").font = TITLE_FONT
        _write_xlsx_blocks(ws2, blocks_translated, start_row=3)
        _size_columns(ws2, _max_table_cols(blocks_translated))

    buf = io.BytesIO()
    wb.save(buf)
    return _stream(buf, fname, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── JSONL — one self-contained record per paragraph/heading/table row,
#            matching the guide's "sentence, not a raw table fragment"
#            principle for rows specifically. Download format only — the
#            Knowledge Base's own ingest endpoint always re-chunks raw
#            text itself (see chatbot.py), this isn't a new ingest path. ──

def _row_sentence(header: list, row: list) -> str:
    return "; ".join(f"{h}: {v}" for h, v in zip(header, row) if v)


def render_jsonl_response(blocks: list, blocks_translated: list | None, doc_meta: dict, fname: str) -> StreamingResponse:
    translated_by_index = {}
    if blocks_translated and len(blocks_translated) == len(blocks):
        translated_by_index = dict(enumerate(blocks_translated))

    lines = []
    idx = 0
    for i, b in enumerate(blocks):
        tb = translated_by_index.get(i)
        if b.get("type") in ("heading", "paragraph"):
            if not b.get("text"):
                continue
            lines.append(json.dumps({
                "id": f"{doc_meta.get('job_id')}-{idx:04d}",
                "text": (tb or {}).get("text") or b["text"],
                "text_original": b["text"],
                "block_type": b["type"],
                "metadata": {**doc_meta, "block_index": i},
            }, ensure_ascii=False))
            idx += 1
        elif b.get("type") == "table" and b.get("rows"):
            header, *body = b["rows"]
            t_header, t_body = None, []
            if tb and tb.get("rows"):
                t_header, *t_body = tb["rows"]
            for ri, row in enumerate(body):
                sentence = _row_sentence(header, row)
                if not sentence:
                    continue
                translated_sentence = _row_sentence(t_header or header, t_body[ri]) if ri < len(t_body) else sentence
                lines.append(json.dumps({
                    "id": f"{doc_meta.get('job_id')}-{idx:04d}",
                    "text": translated_sentence,
                    "text_original": sentence,
                    "block_type": "table_row",
                    "metadata": {**doc_meta, "block_index": i, "row_index": ri},
                }, ensure_ascii=False))
                idx += 1

    buf = io.BytesIO(("\n".join(lines) + ("\n" if lines else "")).encode("utf-8"))
    return _stream(buf, fname, "application/x-ndjson")
